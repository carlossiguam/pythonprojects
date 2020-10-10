from docx import Document
from docx.shared import Inches

#=======================================================
# GLOBALS
#=======================================================
v_NUT = 'EMOV EP-2019-2268'
v_espacio = ' '
v_mesdepago = 'noviembre'
v_objetocontrato = 'RENOVACION DE ENLACES INTERURBANOS PARA LA EMOV EP'
v_contratosercop = 'RE-056-EMOV EP-2018'


document = Document()

document.add_picture('images\logoemovalcaldia.png', width=Inches(6))

document.add_heading('INFORME TECNICO PREVIO PAGO DE FACTURA', level=0)

document.add_heading('DESCRIPCION DEL SERVICIO', level=1)
p = document.add_paragraph('Este informe técnico pretende establecer el pago del servicio por el mes de ')
p.add_run(v_mesdepago).bold = True
p.add_run(' correspondiente al tramite quipux NUT')
p.add_run(v_espacio)
p.add_run(v_NUT).bold = True
p.add_run(' cuyo objeto de contrato es "')
p.add_run(v_objetocontrato).bold = True
p.add_run(', entregado por el proveedor CNT (CORPORACION NACIONAL DE TELECOMUNICACIONES)')
p.add_run(' , dentro del proceso de régimen especial')
p.add_run(v_espacio)
p.add_run(v_contratosercop).bold = True
p.add_run(' , contrato a ejecutarse por 2 años y que fenece el 13 de diciembre de 2020.')


document.add_heading('DETALLE DEL SERVICIO', level=1)

pdetailservice = document.add_paragraph('Este servicio comprende la provisión de enlaces \
    de datos para establecer la comunicación entre la EMOV EP y la ANT, \
    así como los servicios móviles avanzados, de los cuales hace uso por parte de los \
    funcionarios de RTV, cuando estos realizan sus operativos y acceden al sistema de \
    RTV propio de la empresa para ejecutar sus trabajos, desde cualquier punto de la ciudad.')

document.add_heading('OBSERVACIONES', level=1)
pObservaciones = document.add_paragraph('Se informa que la empresa CNT, del mes de ')
pObservaciones.add_run(v_mesdepago).bold = True
pObservaciones.add_run(' ha cumplido satisfactoriamente con la entrega del servicio según la tabla que se indica \
    a continuación que muestra tanto el detalle de los enlaces de datos ANT y SERVICIOS DE RED MOVIL AVANZADO:')




document.add_heading('SOLICITUD DE PAGO DEL SERVICIO  Y PARTIDAS INVOLUCRADAS', level=1)

pmensajefinal = document.add_paragraph('Por lo anteriormente expuesto, se solicita al\
    departamento financiero ejecute el pago correspondiente al mes de septiembre de este\
    servicio revisando las facturas adjuntas a este documento.')

document.add_heading('RESPONSABILIDADES TECNICAS DEL INFORME', level=1)




document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)



records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('INFOTEC PAGO FACTURA OCT 2020.docx')